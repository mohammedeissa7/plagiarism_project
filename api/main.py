from fastapi import FastAPI
from pydantic import BaseModel
from typing import List
import requests
import io
import numpy as np
from docx import Document
from sentence_transformers import SentenceTransformer

app = FastAPI()

model = SentenceTransformer("all-MiniLM-L6-v2")


class UrlRequest(BaseModel):
    urls: List[str]


def extract_docx_text(file_bytes):
    doc = Document(io.BytesIO(file_bytes))
    text = []
    
    for p in doc.paragraphs:
        if p.text.strip():
            text.append(p.text.strip())

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text.append(cell.text.strip())

    return "\n".join(text)


def text_to_embedding(text):
    chunks = [c for c in text.split("\n") if c.strip()]
    emb = model.encode(chunks)
    return np.mean(emb, axis=0)


def similarity_matrix(texts):

    embeddings = [text_to_embedding(t) for t in texts]

    embeddings = np.array(embeddings)

    norms = np.linalg.norm(embeddings, axis=1, keepdims=True)
    normalized = embeddings / norms

    sim = np.dot(normalized, normalized.T)

    return (sim * 100).round(2).tolist()


@app.post("/similarity/from-urls")
async def similarity_from_urls(req: UrlRequest):

    texts = []
    filenames = []

    for url in req.urls:

        r = requests.get(url)

        if r.status_code != 200:
            raise Exception(f"Cannot download file: {url}")

        text = extract_docx_text(r.content)

        texts.append(text)

        filenames.append(url.split("/")[-1])

    matrix = similarity_matrix(texts)

    return {
        "files": filenames,
        "matrix": matrix
    }