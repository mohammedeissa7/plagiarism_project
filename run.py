
import os
import sys
import argparse
import pandas as pd
import numpy as np
from docx import Document
from sentence_transformers import SentenceTransformer, util


print("Current working directory:", os.getcwd())

class PlagiarismChecker:
    def __init__(self, model_name="all-MiniLM-L6-v2", device=None):
        try:
            # device=None lets SentenceTransformer pick (cuda if available)
            self.model = SentenceTransformer(model_name, device=device)
        except Exception as e:
            print("[!] Failed to load model:", e)
            print("    - Check internet connection (first run needs to download model) or model name.")
            raise

    def read_docx(self, path):
        """Extract text from paragraphs and tables (fallback to empty string)."""
        doc = Document(path)
        parts = []
        
        for p in doc.paragraphs:
            if p.text and p.text.strip():
                parts.append(p.text.strip())
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    txt = cell.text.strip()
                    if txt:
                        parts.append(txt)
        
        return "\n".join(parts)

    def load_documents(self, folder_path):
        """Load all .docx files from a folder"""
        if not os.path.isdir(folder_path):
            raise FileNotFoundError(f"Folder not found: {folder_path}")

        files = [f for f in os.listdir(folder_path) if f.lower().endswith(".docx")]
        files.sort()
        if not files:
            raise FileNotFoundError(f"No .docx files found in folder: {folder_path}")

        texts = []
        for f in files:
            full_path = os.path.join(folder_path, f)
            texts.append(self.read_docx(full_path))

        return files, texts

    def _doc_to_embedding(self, text, chunk_by="paragraph"):

        if not text or not text.strip():
            # create zero vector with same dimension as model (encode a dummy to infer dim)
            dummy = self.model.encode([""], convert_to_tensor=True)
            return dummy[0].cpu().numpy() * 0.0

        # split into chunks (we expect paragraphs separated by newlines)
        chunks = [c.strip() for c in text.split("\n") if c.strip()]
        # encode chunks
        chunk_embs = self.model.encode(chunks, convert_to_tensor=True)
        # average chunk embeddings -> document embedding
        doc_emb = chunk_embs.mean(dim=0)
        return doc_emb.cpu().numpy()

    def compute_similarity_matrix(self, texts):
        
        if not texts:
            raise ValueError("Please provide at least some text to compare.")

        
        embeddings = []
        for t in texts:
            emb = self._doc_to_embedding(t)
            embeddings.append(emb)
        embeddings = np.vstack(embeddings)           

       
        # normalize
        norms = np.linalg.norm(embeddings, axis=1, keepdims=True)
        # avoid divide-by-zero
        norms[norms == 0] = 1.0
        normalized = embeddings / norms
        cosine_sim = normalized @ normalized.T    

       
        percent_matrix = np.clip(cosine_sim, 0.0, 1.0) * 100.0
        percent_matrix = np.round(percent_matrix, 2)
        return percent_matrix

    def save_matrix_csv(self, files, matrix, output_name="similarity_matrix.csv"):
        df = pd.DataFrame(matrix, index=files, columns=files)
        df.to_csv(output_name)
        print(f"[+] Saved: {output_name}")
        return df


def main():
    parser = argparse.ArgumentParser(description="Simple docx plagiarism similarity checker")
    parser.add_argument("--folder", "-f", default="students_docs", help="Folder containing .docx files")
    parser.add_argument("--output", "-o", default="similarity_matrix.csv", help="Output CSV file")
    parser.add_argument("--model", "-m", default="all-MiniLM-L6-v2", help="SentenceTransformer model name")
    args = parser.parse_args()

    try:
        checker = PlagiarismChecker(model_name=args.model)
    except Exception:
        print("Model load failed; exiting.")
        sys.exit(1)

    try:
        files, texts = checker.load_documents(args.folder)
    except Exception as e:
        print("[!] Error loading documents:", e)
        sys.exit(1)

    matrix = checker.compute_similarity_matrix(texts)
    df = checker.save_matrix_csv(files, matrix, args.output)

    print("\nSimilarity Matrix (%):")
    print(df)


if __name__ == "__main__":
    main()
